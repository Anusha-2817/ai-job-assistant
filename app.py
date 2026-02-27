import json
import streamlit as st
from main import coordinator
from agents import SkillOntologyAgent

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
from datetime import datetime

with open("data.json", "r") as f:
    ontology_data = json.load(f)

ontology_agent = SkillOntologyAgent(ontology_data)

def generate_docx(text, user_name):
    document = Document()

    # Set margins
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add Name (Bold, Larger Font)
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(user_name)
    name_run.bold = True
    name_run.font.name = "Times New Roman"
    name_run.font.size = Pt(14)

    # Add Date (Right aligned)
    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run(datetime.today().strftime("%B %d, %Y"))
    date_run.font.name = "Times New Roman"
    date_run.font.size = Pt(12)

    # Add spacing
    document.add_paragraph("")

    # Add cover letter content
    for paragraph in text.split("\n\n"):
        p = document.add_paragraph(paragraph)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


st.set_page_config( #layout
    page_title="AI Resume Analyzer",
    layout="wide"
)

st.title("🤖 AI Resume Analyzer")
st.write("Paste a job description and your resume to see how well you match.")
user_name = st.text_input("Your Full Name (for Cover Letter)")
col1, col2 = st.columns(2)

with col1:
    job_text = st.text_area(
        "📌 Job Description",
        height=250,
        placeholder="Paste the job description here..."
    )

with col2:
    resume_text = st.text_area(
        "📄 Resume",
        height=250,
        placeholder="Paste your resume here..."
    )

if st.button("🔍 Analyze"):
    if not job_text or not resume_text:
        st.warning("Please enter both job description and resume.")
    else:
        with st.spinner("Analyzing..."):
            result = coordinator(job_text, resume_text, ontology_agent, user_name)

        st.success("Done!")

        match_percent = result["skill_match"]["match_percentage"]
        match_percent = round(match_percent, 1)
        st.metric(
            label="Resume Match Percentage",
            value=f"{match_percent:.2f}%"
        )
        st.subheader("🧠 Extracted Skills")

        c1, c2 = st.columns(2)

        with c1:
            st.markdown("### Job Skills")
            for skill in result["job_skills"]:
                st.markdown(f"- **{skill}**")

        with c2:
            st.markdown("### Resume Skills")
            for skill in result["resume_skills"]:
                st.markdown(f"- **{skill}**")

        st.subheader("🧠 Ontology-Enriched Skills")

        c3, c4 = st.columns(2)
        with c3:
            st.markdown(f"### Enriched Job Skills ({len(result['enriched_job_skills'])})")
            for skill in result["enriched_job_skills"]:
               st.markdown(
                    f"- **{skill['skill']}**  \n"
                    f"    - *Category:* {skill['category']}"
                )
        with c4:
            st.markdown(f"### Enriched Resume Skills ({len(result['enriched_resume_skills'])})")
            for skill in result["enriched_resume_skills"]:
                st.markdown(
                    f"- **{skill['skill']}**  \n"
                    f"    - *Category:* {skill['category']}"
                )

        st.subheader("📊 Skill Match Analysis")

        match_percent = result["skill_match"]["match_percentage"]
        missing_skills = result["skill_match"]["missing_skills"]
        enriched_job = result["enriched_job_skills"]
        # st.metric("Match Percentage", f"{match_percent}%")
        c5, c6 = st.columns(2)

        with c5:
            st.markdown("###Matched Skills")
            for skill in result["skill_match"]["matched_skills"]:
                st.markdown(f"- {skill}")
        with c6:
            st.markdown("###Missing Skills")
            if not missing_skills:
                    st.success("No missing skills 🎉 Great match!")
            for skill in result["skill_match"]["missing_skills"]:
                st.markdown(f"- {skill}")
            for skill in enriched_job:

                if skill["skill"] in missing_skills:
                    st.markdown(
                        f"- **{skill['skill']}**"
                        f"  - Category: {skill['category']}"
                    )
                
        if match_percent >= 75:
            st.success("Strong Match 🚀")
        elif match_percent >= 50:
            st.warning("Moderate Match ⚡")
        else:
            st.error("Low Match ❗ Improve Resume")

        def format_cover_letter(text):
            paragraphs = text.split("\n")
            clean_paragraphs = [p.strip() for p in paragraphs if p.strip()]
            return "\n\n".join(clean_paragraphs)

        formatted_letter = format_cover_letter(result["cover_letter"])


        st.subheader("✍️ AI-Generated Cover Letter")

        st.markdown(
            f"""
            <div style='
                background-color:#ffffff;
                padding:25px;
                border-radius:12px;
                border:1px solid #e6e6e6;
                box-shadow: 0 2px 8px rgba(0,0,0,0.05);
                font-family: Times New Roman;
                font-size:16px;
                line-height:1.4;
                color:#1a1a1a;;
            '>
            {formatted_letter.replace('\n\n', '<br><br>').replace('\n', '<br>')}
            </div>
            """,
            unsafe_allow_html=True
        )
        st.markdown("<hr>", unsafe_allow_html=True)
        docx_file = generate_docx(formatted_letter, user_name)

        st.download_button(
            label="📄 Download Professional Cover Letter (.docx)",
            data=docx_file,
            file_name="Professional_Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
       
